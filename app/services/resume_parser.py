"""
Resume parsing service for extracting candidate information from resumes
"""
import os
import re
import json
try:
    import spacy
except ImportError:  # pragma: no cover - optional dependency
    spacy = None

try:
    import nltk
except ImportError:  # pragma: no cover - optional dependency
    nltk = None

try:
    import pandas as pd
except ImportError:  # pragma: no cover - optional dependency
    pd = None
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path
import mimetypes 
try:
    from pdfminer.high_level import extract_text as extract_pdf_text
except ImportError:  # pragma: no cover - optional dependency
    extract_pdf_text = None

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None
import io
try:
    from nltk.tokenize import word_tokenize
    from nltk.corpus import stopwords
except Exception:  # pragma: no cover - optional dependency
    word_tokenize = None
    stopwords = None
import logging

# Download NLTK data when available
if nltk is not None:
    try:
        nltk.download('punkt', quiet=True)
        nltk.download('stopwords', quiet=True)
        nltk.download('averaged_perceptron_tagger', quiet=True)
    except Exception:
        pass

logger = logging.getLogger(__name__)

class ResumeParser:
    """Resume parsing service to extract candidate information"""
    
    def __init__(self):
        """Initialize resume parser with NLP models"""
        self.nlp = None
        if spacy is not None:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except Exception:
                # Keep parser functional without spaCy model.
                self.nlp = None
        
        # Common patterns for extraction
        self.email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        self.phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        self.linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/(?:in|company)/[a-zA-Z0-9_-]+'
        
        # Skill keywords database
        self.skill_keywords = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'node.js', 'express'],
            'database': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ci/cd'],
            'data_science': ['pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn', 'machine learning', 'ai'],
            'devops': ['jenkins', 'git', 'github', 'gitlab', 'ansible', 'prometheus', 'grafana'],
            'soft_skills': ['communication', 'leadership', 'teamwork', 'problem solving', 'time management']
        }
        
    class ResumeParser:
        
        SKILL_CATEGORIES = {
        "cloud": [...],
        "data_science": [...],
        "devops": [...],
        "soft_skills": [...]
    }

    def extract_text_from_file(self, file_path: str) -> str:
        mime_type, _ = mimetypes.guess_type(file_path)

        if mime_type is None:
            raise ValueError("Could not determine file type")

        if mime_type == "application/pdf":
            return self._extract_from_pdf(file_path)

        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return self._extract_from_docx(file_path)

        elif mime_type.startswith("text/"):
            return self._extract_from_text(file_path)

        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        if extract_pdf_text is None:
            logger.warning("pdfminer is not installed; PDF parsing is unavailable.")
            return ""
        try:
            text = extract_pdf_text(file_path)
            # Clean up text
            text = re.sub(r'\s+', ' ', text)
            text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Remove non-ASCII
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        if Document is None:
            logger.warning("python-docx is not installed; DOCX parsing is unavailable.")
            return ""
        try:
            doc = Document(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += '\n' + cell.text
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except:
            with open(file_path, 'r', encoding='latin-1', errors='ignore') as f:
                return f.read().strip()
    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """
        Parse resume and extract structured information
        
        Returns:
            Dictionary with extracted information
        """
        try:
            # Extract raw text
            raw_text = self.extract_text_from_file(file_path)
            
            if not raw_text:
                return {"error": "Could not extract text from resume"}
            
            # Extract various information
            result = {
                "raw_text": raw_text[:500] + "...",  # Store first 500 chars
                "name": self._extract_name(raw_text),
                "email": self._extract_email(raw_text),
                "phone": self._extract_phone(raw_text),
                "linkedin": self._extract_linkedin(raw_text),
                "skills": self._extract_skills(raw_text),
                "experience": self._extract_experience(raw_text),
                "education": self._extract_education(raw_text),
                "summary": self._extract_summary(raw_text),
                "companies": self._extract_companies(raw_text),
                "locations": self._extract_locations(raw_text)
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            return {"error": str(e)}
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume"""
        # Look for name patterns at the beginning of the resume
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and len(line) < 50:  # Name is usually short
                # Check if line looks like a name (contains letters and maybe spaces)
                if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', line):
                    return line
        
        # Fallback: Use the first line that's not empty and not too long
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) < 100:
                return line
        
        return ""
    
    def _extract_email(self, text: str) -> str:
        """Extract email address"""
        emails = re.findall(self.email_pattern, text)
        return emails[0] if emails else ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number"""
        phones = re.findall(self.phone_pattern, text)
        # Clean phone number
        if phones:
            phone = re.sub(r'[^\d+]', '', phones[0])
            return phone[:15]  # Limit length
        return ""
    
    def _extract_linkedin(self, text: str) -> str:
        """Extract LinkedIn profile"""
        linkedin = re.findall(self.linkedin_pattern, text, re.IGNORECASE)
        return linkedin[0] if linkedin else ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        skills = []
        text_lower = text.lower()
        
        # Check for each skill category
        for category, skill_list in self.skill_keywords.items():
            for skill in skill_list:
                if skill.lower() in text_lower:
                    skills.append(skill)
        
        # Also look for skills mentioned in a skills section
        skills_section_patterns = [
            r'skills(?:\s*[:•])(.*?)(?:\n\n|\Z)',
            r'technical skills(?:\s*[:•])(.*?)(?:\n\n|\Z)',
            r'core competencies(?:\s*[:•])(.*?)(?:\n\n|\Z)'
        ]
        
        for pattern in skills_section_patterns:
            match = re.search(pattern, text_lower, re.IGNORECASE | re.DOTALL)
            if match:
                section_text = match.group(1)
                # Extract individual skills from the section
                section_skills = re.findall(r'[a-zA-Z0-9+#\.\s]+', section_text)
                for skill in section_skills:
                    skill = skill.strip()
                    if skill and len(skill) < 50:
                        skills.append(skill)
        
        # Remove duplicates and clean up
        skills = list(set([s.strip().title() for s in skills if s.strip()]))
        return skills[:20]  # Limit to 20 skills
    
    def _extract_experience(self, text: str) -> int:
        """Extract total years of experience"""
        # Look for patterns like "5 years", "10+ years", etc.
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of)?\s*experience',
            r'experience\s*:\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*in\s*[a-zA-Z\s]+'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    years = int(match)
                    if 0 <= years <= 50:  # Reasonable range
                        return years
                except:
                    pass
        
        # Try to calculate from employment history
        # This is a simplified calculation
        experience_keywords = ['years', 'yrs', 'experience']
        lines = text.lower().split('\n')
        for line in lines:
            for keyword in experience_keywords:
                if keyword in line:
                    # Look for numbers near the keyword
                    numbers = re.findall(r'\b(\d+)\b', line)
                    for num in numbers:
                        try:
                            years = int(num)
                            if 1 <= years <= 50:
                                return years
                        except:
                            pass
        
        return 0
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        
        # Common education keywords
        education_keywords = [
            'university', 'college', 'institute', 'school',
            'bachelor', 'master', 'phd', 'mba', 'bsc', 'msc',
            'degree', 'diploma', 'certification'
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()
            # Check if line contains education keywords
            if any(keyword in line_lower for keyword in education_keywords):
                edu_entry = {
                    "institution": line.strip(),
                    "degree": self._extract_degree(line),
                    "year": self._extract_year(line)
                }
                education.append(edu_entry)
        
        return education[:5]  # Limit to 5 education entries
    
    def _extract_degree(self, text: str) -> str:
        """Extract degree from education line"""
        degrees = ['B.Sc', 'B.Tech', 'B.E', 'B.A', 'M.Sc', 'M.Tech', 'M.A', 'PhD', 'MBA']
        for degree in degrees:
            if degree.lower() in text.lower():
                return degree
        return ""
    
    def _extract_year(self, text: str) -> str:
        """Extract year from education line"""
        years = re.findall(r'(19|20)\d{2}', text)
        return years[0] if years else ""
    
    def _extract_summary(self, text: str) -> str:
        """Extract summary/objective from resume"""
        summary_patterns = [
            r'summary(?:\s*[:•])(.*?)(?:\n\n|\Z)',
            r'objective(?:\s*[:•])(.*?)(?:\n\n|\Z)',
            r'profile(?:\s*[:•])(.*?)(?:\n\n|\Z)'
        ]
        
        for pattern in summary_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip()
                # Clean up summary
                summary = re.sub(r'\s+', ' ', summary)
                return summary[:500]  # Limit length
        
        # Fallback: Use first paragraph
        paragraphs = text.split('\n\n')
        if paragraphs:
            return paragraphs[0][:300]
        
        return ""
    
    def _extract_companies(self, text: str) -> List[str]:
        """Extract company names from work experience"""
        companies = []
        
        # Common company indicators
        experience_patterns = [
            r'at\s+([A-Z][a-zA-Z\s&\.]+)(?:\s+|$)',
            r'company\s*:\s*([A-Z][a-zA-Z\s&\.]+)',
            r'employer\s*:\s*([A-Z][a-zA-Z\s&\.]+)'
        ]
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, text)
            companies.extend(matches)
        
        # Look for company names in all caps or with Inc/LLC/etc
        company_pattern = r'\b([A-Z][A-Za-z\s&\.]+(?:Inc|LLC|Ltd|Corp|Corporation|Company))\b'
        matches = re.findall(company_pattern, text)
        companies.extend(matches)
        
        # Clean and deduplicate
        companies = list(set([c.strip() for c in companies if len(c) > 2]))
        return companies[:10]
    
    def _extract_locations(self, text: str) -> List[str]:
        """Extract locations from resume"""
        # This is a simplified location extraction
        # In production, use a proper NER model or geocoding service
        location_keywords = ['location', 'address', 'city', 'based in']
        locations = []
        
        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in location_keywords):
                # Extract potential location
                location_match = re.search(r':\s*(.+)$', line)
                if location_match:
                    locations.append(location_match.group(1).strip())
        
        return locations

# Singleton instance
resume_parser = ResumeParser()
